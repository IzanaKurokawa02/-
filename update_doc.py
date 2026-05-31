import re
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def apply_gost(doc):
    """Применить параметры ГОСТа ко всем страницам"""
    for section in doc.sections:
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(1.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
    
    for paragraph in doc.paragraphs:
        p = paragraph.paragraph_format
        p.line_spacing = 1.5
        p.first_line_indent = Cm(1.25)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)

def fix_title_page(doc):
    """Добавить соавторов и группу на титульный лист"""
    for para in doc.paragraphs:
        if "Выполнил:" in para.text and "25132" not in para.text:
            para.text = "Выполнили: обучающиеся группы 25132БУ\nБельтиков Максим, Прохоров Артём, Черняков Кирилл, Воробьёв Мирослав"
            para.style = doc.styles['Normal']
            break

def fix_survey_data(doc):
    """Заменить старые цифры опроса (86 респондентов) на новые (104)"""
    # Замены для общего числа респондентов
    for para in doc.paragraphs:
        if '86 респондентов' in para.text:
            para.text = para.text.replace('86 респондентов', '104 респондента')
        if '86 респ.' in para.text:
            para.text = para.text.replace('86 респ.', '104 респ.')

    # Список замен (старая строка -> новая)
    replacements = [
        ('54,7%', '58,7%'), ('16,3%', '13,5%'), ('9,3%', '9,6%'), ('8,1%', '6,7%'),
        ('7,0%', '5,8%'), ('4,7%', '3,8%'), ('2,3%', '1,9%'), ('42,1%', '45,6%'),
        ('27,4%', '23,7%'), ('20%', '21,1%'), ('47,7%', '49%'), ('43%', '44,2%'),
        ('17,4%', '15,4%'), ('81,4%', '82,7%'), ('44,2%', '42,3%'), ('37,2%', '40,4%'),
        ('15,1%', '14,4%'), ('25,1%', '24,2%'), ('17,3%', '16,3%'), ('12,6%', '13,2%'),
        ('12,0%', '11,9%'), ('11,5%', '11,9%'), ('11,0%', '11,5%'), ('9,4%', '9,7%'),
        ('38,4%', '38,5%'), ('31,4%', '32,7%'), ('47,7%', '50%'), ('18,6%', '19,2%'),
        ('10,5%', '9,6%'), ('8,1%', '8,7%'), ('36,0%', '33,7%'), ('52,3%', '55,8%'),
        ('11,6%', '10,6%'), ('16,7%', '16,3%'), ('15,1%', '14,8%'), ('12,2%', '11,9%'),
        ('8,2%', '9%'), ('6,6%', '7,3%'), ('6,1%', '5,6%'), ('3,4%', '4,2%')
    ]
    for old, new in replacements:
        for para in doc.paragraphs:
            if old in para.text:
                para.text = para.text.replace(old, new)

    # Дополнительно ищем отдельное число 86
    for para in doc.paragraphs:
        para.text = re.sub(r'\b86\b', '104', para.text)

def main():
    input_file = "Document 4.docx"
    output_file = "ГОТОВЫЙ_ПРОЕКТ.docx"
    
    doc = Document(input_file)
    apply_gost(doc)
    fix_title_page(doc)
    fix_survey_data(doc)
    
    # Изменить заголовок файла (свойства)
    doc.core_properties.title = "Глава 2. Исследование потребностей пользователей и проектирование агрегатора – Агрегатор услуг как новое окно интернет-активности пользователя"
    
    doc.save(output_file)
    print(f"✅ Готово! Файл сохранён как {output_file}")

if __name__ == "__main__":
    main()
